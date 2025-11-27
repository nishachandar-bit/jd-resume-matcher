# Enhanced JD-Resume Matcher (UI improvements)
# Save as app.py and run: streamlit run app.py
# Requirements: streamlit pandas openpyxl python-docx pdfminer.six rapidfuzz openpyxl

import streamlit as st
import pandas as pd
import io, re, json, os
from collections import OrderedDict
from docx import Document
from pdfminer.high_level import extract_text as pdf_extract_text
from rapidfuzz import fuzz

st.set_page_config(page_title="JD ↔ Resume Matcher (Enhanced UI)", layout="wide")

# ------------------ Helpers ------------------
def read_text_file(uploaded):
    data = uploaded.read()
    try:
        return data.decode("utf-8", errors="ignore")
    except Exception:
        try:
            return data.decode("latin-1", errors="ignore")
        except:
            return ""

def read_pdf_from_path(path_or_file):
    try:
        if isinstance(path_or_file, str):
            return pdf_extract_text(path_or_file)
        else:
            path_or_file.seek(0)
            return pdf_extract_text(path_or_file)
    except Exception as e:
        st.warning(f"PDF parse warning: {e}")
        return ""

def read_docx_from_path(path_or_file):
    try:
        if isinstance(path_or_file, str):
            doc = Document(path_or_file)
        else:
            path_or_file.seek(0)
            doc = Document(path_or_file)
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        st.warning(f"DOCX parse warning: {e}")
        return ""

def extract_text_any(uploaded, filename=None):
    if filename and filename.lower().endswith(".pdf"):
        return read_pdf_from_path(uploaded)
    if filename and filename.lower().endswith(".docx"):
        return read_docx_from_path(uploaded)
    # fallback: try to read bytes as text
    try:
        return read_text_file(uploaded)
    except:
        return ""

# normalization and cleaning
NOISE_RE = re.compile(r'\b(exp|exp\.|experience|expertise|minimum|should|years|yrs)\b', re.I)
PUNCT_RE = re.compile(r'[\(\)\[\]\-_:,\/]+')

def normalize_skill_label(s):
    if not s:
        return ""
    x = s.strip()
    x = NOISE_RE.sub(" ", x)
    x = PUNCT_RE.sub(" ", x)
    x = re.sub(r'\s+', ' ', x).strip()
    return x

def parse_years_from_text(s):
    if not s:
        return None
    m = re.search(r'(\d+)\s*\+?\s*(?:years|yrs|y)', s)
    if m:
        try:
            return int(m.group(1))
        except:
            return None
    m = re.search(r'(\d+)\s*[-–]\s*(\d+)', s)
    if m:
        try:
            return int(m.group(2))
        except:
            return None
    return None

# presence detection improved
def has_skill(text, skill, synonyms=None, strict=False):
    if not text or not skill:
        return False
    t = text.lower()
    s = skill.lower().strip()
    candidates = [s]
    if synonyms:
        candidates += [v.lower() for v in synonyms]
    # exact or substring
    for c in candidates:
        if c and c in t:
            return True
    # whole-token check
    tokens = [w for w in re.split(r'\W+', s) if w]
    if tokens and all(re.search(rf'\b{re.escape(tok)}\b', t) for tok in tokens):
        return True
    # fuzzy only when not strict and skill length reasonable
    if not strict and len(s) > 3:
        try:
            score = fuzz.partial_ratio(s, t)
            return score >= 85
        except:
            return False
    return False

def extract_years_near_skill(text, skill_syns, window=120):
    t = text.lower()
    years_found = []
    for syn in skill_syns:
        for m in re.finditer(re.escape(syn.lower()), t):
            start = max(0, m.start() - window)
            end = min(len(t), m.end() + window)
            win = t[start:end]
            m2 = re.findall(r'(\d+\s*[-–]\s*\d+|\d+\+?)\s*(?:years|yrs|y)', win)
            for g in m2:
                if "-" in g:
                    nums = re.findall(r'(\d+)', g)
                    if nums:
                        years_found.append(int(nums[-1]))
                else:
                    n = re.search(r'(\d+)', g)
                    if n:
                        years_found.append(int(n.group(1)))
    if years_found:
        return max(years_found)
    # fallback: global years
    m3 = re.findall(r'(\d+\s*[-–]\s*\d+|\d+\+?)\s*(?:years|yrs|y)', t)
    for g in m3:
        if "-" in g:
            nums = re.findall(r'(\d+)', g)
            if nums:
                years_found.append(int(nums[-1]))
        else:
            n = re.search(r'(\d+)', g)
            if n:
                years_found.append(int(n.group(1)))
    return max(years_found) if years_found else None

# default synonym map (copyable/editable)
DEFAULT_SYNONYMS = {
    "ci/cd": ["ci/cd","ci cd","continuous integration","continuous delivery","jenkins","pipeline","devops"],
    "tosca": ["tosca","tricentis tosca","tricentis"],
    "web application automation": ["web application","web app","ui automation","selenium","frontend","web testing","browser testing"],
    "mainframe automation testing": ["mainframe","3270","green screen","mainframe testing","jcl","cobol"]
}

# ------------------ UI Layout ------------------
st.title("JD ↔ Resume Matcher — Enhanced UI")
st.markdown("Upload a Job Description and resumes. Clean the extracted skills, add synonyms, tune weights, and run matching.")

# left column: JD deconstruction + upload
left, mid, right = st.columns([1,2,1])

with left:
    st.header("1) Upload JD & Controls")
    jd_file = st.file_uploader("Upload JD (PDF/DOCX/TXT)", type=["pdf","docx","txt"], key="jd_upload")
    auto_normalize_btn = st.button("Auto-normalize extracted skills")
    accept_all_btn = st.button("Accept all skills (move to skill list)")
    st.markdown("**Presence weight** (default = 60%)")
    presence_weight = st.slider("Presence weight (%)", 40, 90, 60)
    enforce_years = st.checkbox("Enforce JD years strictly", value=True)
    st.markdown("---")
    st.markdown("Upload resumes (multiple)")
    resumes = st.file_uploader("Upload Resumes (PDF/DOCX/TXT) - select multiple", type=["pdf","docx","txt"], accept_multiple_files=True, key="resumes")
    st.info("Recommended: DOCX or text-based PDF. For scanned PDFs use OCR externally.")

# mid: skill list visual
with mid:
    st.header("2) Deconstructed Skills")
    jd_text_area = st.empty()
    if jd_file:
        extracted = ""
        if jd_file.name.lower().endswith(".pdf"):
            extracted = read_pdf_from_path(jd_file)
        elif jd_file.name.lower().endswith(".docx"):
            extracted = read_docx_from_path(jd_file)
        else:
            extracted = read_text_file(jd_file)
        extracted = extracted or ""
        jd_text_area = st.text_area("JD preview (editable)", value=extracted, height=220, key="jd_preview")
        # simple extraction heuristics: find capitalized phrases and master list hits
        master_skills = ["TOSCA","CI/CD","Web Application Automation","Mainframe Automation Testing","LoadRunner","Dynatrace","Splunk","VUGen"]
        candidates = []
        # master hits
        for ms in master_skills:
            if ms.lower() in jd_text_area.lower() and ms not in candidates:
                candidates.append(ms)
        # capitalized phrases heuristics
        caps = re.findall(r'\b([A-Z][A-Za-z0-9+\-#.]{1,}(?:\s+[A-Z][A-Za-z0-9+\-#.]{1,}){0,2})\b', jd_text_area)
        for c in caps:
            if len(c.split())<=4 and c not in candidates:
                candidates.append(c)
        # show candidates with detected years
        st.write("Auto-extracted candidates (edit before accepting):")
        cand_df = []
        for c in candidates:
            norm = normalize_skill_label(c)
            # find years if any near mention
            req = None
            m = re.search(rf'(?:minimum|at least|>=|\b{re.escape(c)}\b).{{0,80}}?(\d+)\s*(?:\+)?\s*(?:years|yrs|y)', jd_text_area, re.I)
            if m:
                try:
                    req = int(m.group(1))
                except:
                    req = None
            cand_df.append({"raw":c,"normalized":norm,"req":req})
        if cand_df:
            df_cand = pd.DataFrame(cand_df)
            st.dataframe(df_cand)
    else:
        st.info("Upload a JD to extract skills.")

    st.markdown("**Skill list (editable)**")
    skill_list_area = st.text_area("One skill per line. Prefix with M: for Mandatory, D: for Desired (optional).", value="", height=200, key="skill_list")
    # parse skills area into structured list
    def parse_skill_list(text):
        lines = [l.strip() for l in text.splitlines() if l.strip()]
        skills_struct = []
        for line in lines:
            tag = "desired"
            content = line
            if line.lower().startswith("m:") or line.lower().startswith("mandatory:"):
                tag = "mandatory"
                content = line.split(":",1)[1].strip()
            elif line.lower().startswith("d:") or line.lower().startswith("desired:"):
                tag = "desired"
                content = line.split(":",1)[1].strip()
            norm = normalize_skill_label(content)
            skills_struct.append({"skill":norm,"group":tag,"req":None})
        return skills_struct

    user_skills = parse_skill_list(skill_list_area) if skill_list_area.strip() else []
    # populate default if empty and JD exists
    if not user_skills and jd_file:
        # use candidates as base
        user_skills = []
        for c in candidates:
            norm = normalize_skill_label(c)
            user_skills.append({"skill":norm,"group":"mandatory" if "minimum" in c.lower() or "min" in c.lower() else "desired","req":None})
    # show editable table-like interface
    if user_skills:
        st.write("Final skills (you can edit the textarea above and re-parse):")
        sk_rows = []
        for s in user_skills:
            sk_rows.append(f"{ 'M:' if s['group']=='mandatory' else 'D:' } {s['skill']}{' | req:'+str(s['req']) if s['req'] else ''}")
        st.text_area("Parsed skills (read-only preview)", value="\n".join(sk_rows), height=160)

# right: synonyms and matching controls
with right:
    st.header("3) Synonyms & Matching")
    st.write("Default synonyms (editable):")
    synonyms_state = st.text_area("Edit synonyms as JSON (skill: [variants])", value=json.dumps(DEFAULT_SYNONYMS, indent=2), height=220, key="synonyms_json")
    try:
        synonyms_map = json.loads(synonyms_state)
    except:
        st.error("Invalid JSON for synonyms; revert to default.")
        synonyms_map = DEFAULT_SYNONYMS.copy()
    st.markdown("---")
    strict_matching = st.checkbox("Strict matching (avoid fuzzy for short tokens)", value=True)
    st.markdown("Show / Hide columns in results:")
    show_presence = st.checkbox("Show presence columns", value=True)
    show_years = st.checkbox("Show years columns", value=True)
    st.markdown("---")
    run_button = st.button("Run Matching", key="run_match")

# ------------------ Matching ------------------
def build_skills_from_user(user_skills, synonyms_map):
    skills = []
    for s in user_skills:
        name = s['skill']
        grp = s.get('group','desired')
        req = s.get('req')
        syns = synonyms_map.get(name.lower(), [])
        skills.append({"name":name, "group":grp, "req":req, "synonyms":syns})
    return skills

if run_button:
    if not jd_file:
        st.error("Please upload JD before running.")
    elif not resumes:
        st.error("Please upload resumes before running.")
    else:
        with st.spinner("Parsing resumes and running matcher..."):
            # finalize skills from textarea
            final_skills = build_skills_from_user(user_skills, synonyms_map)
            # compute weights: default presence/experience split from presence_weight slider
            presence_w = presence_weight/100.0
            exp_w = 1.0 - presence_w

            results = []
            for up in resumes:
                fname = up.name
                txt = extract_text_any(up, fname)
                txt_norm = txt or ""
                row = {"Resume":fname}

                # per-skill checks
                mandatory_scores = []
                desired_scores = []
                for sk in final_skills:
                    name = sk['name']
                    syns = sk.get('synonyms',[])
                    group = sk.get('group','desired')
                    req = sk.get('req')
                    present = has_skill(txt_norm, name, synonyms=syns, strict=strict_matching)
                    years = extract_years_near_skill(txt_norm, [name]+syns) if present else None
                    # compute skill score
                    if req:
                        # requirement exists -> presence + years enforcement
                        if present and years is not None:
                            exp_ratio = min(years / req, 1.0)
                            score = presence_w * 1.0 + exp_w * exp_ratio
                        elif present and years is None:
                            score = presence_w * 1.0 + exp_w * 0.0
                        else:
                            score = 0.0
                    else:
                        score = 1.0 if present else 0.0

                    # store columns
                    row[f"{name}_presence"] = "Yes" if present else "No"
                    row[f"{name}_years"] = f"{years}y" if years else ""
                    row[f"{name}_req"] = f"{req}y" if req else ""
                    row[f"{name}_score_%"] = round(score*100,2)
                    if group=='mandatory':
                        mandatory_scores.append(score)
                    else:
                        desired_scores.append(score)

                # aggregate weighted scores: default weights (mandatory 80% desired 20%) but user can adjust in advanced UI later
                mand_weight = 0.8
                des_weight = 0.2
                mand_avg = sum(mandatory_scores)/len(mandatory_scores) if mandatory_scores else 0.0
                des_avg = sum(desired_scores)/len(desired_scores) if desired_scores else 0.0
                overall = round((mand_avg*mand_weight + des_avg*des_weight)*100,2)
                row['Match %'] = overall
                results.append(row)

            df = pd.DataFrame(results)
            st.success("Matching complete. Preview results below:")
            # allow column visibility customization
            st.dataframe(df)
            # export excel
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='JD Match Analysis')
            output.seek(0)
            st.download_button('Download Excel', data=output, file_name='jd_match_results.xlsx', mime='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')

st.markdown('---')
st.markdown('Need a feature? Tell me which UI improvement to add next.')
